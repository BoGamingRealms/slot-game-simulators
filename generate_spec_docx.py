import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_game_spec_docx(output_path):
    doc = docx.Document()
    
    # Page setup - Normal Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("GAME SPECIFICATION\nCASH VORTEX: TRIPLE POWER™")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E) # Indigo
    title_p.paragraph_format.space_after = Pt(4)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Game Specification for Frontend Developers, Game Designers & QA | Version 2.0")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub_p.paragraph_format.space_after = Pt(16)
    
    # Metadata Box (Table)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Target Platforms", "Mobile (iOS/Android), Tablet & Desktop (HTML5 / WebGL)"),
        ("Reel Layout", "5×5 Matrix (25 Reel Positions) with Central Wild Star at (2,2)"),
        ("Paylines & Mechanics", "12 Slingo Lines, 3-Spin Coin Lifespans, Lock & Win Respins"),
        ("Key Features", "Center Wild Wheel Bonus, 3-Tier Top X-Wheels, Lock & Slingo™ Bonus")
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = k
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9.5)
        c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        c1.text = v
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c0, "E8EAF6")
        set_cell_background(c1, "F5F5F5")
        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    def add_heading_1(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x79, 0x6B) # Teal
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        return h

    # Section 1
    add_heading_1("1. Executive Summary & High Concept")
    p1 = doc.add_paragraph(
        "Cash Vortex: Triple Power™ combines the excitement of Slingo line completion with persistent locking cash symbols, "
        "explosive modifier mechanics (Strikes and Vortexes), a 3-tier top wheel progression system (X-Wheels), an independent center-reel wheel bonus, "
        "and a dedicated 5×5 Lock & Slingo™ respin bonus game."
    )
    p1.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph(
        "Unlike traditional slots where symbols vanish on each spin, symbols in Cash Vortex hold an active 3-spin lifespan, staying locked on the reels "
        "to help players complete 5-symbol Slingo Lines (horizontal, vertical, and diagonal). When Slingo lines complete, they award the sum of all cash values along that line. "
        "If a completed line passes through the Central Wild Star, it activates the Center Wild Wheel Bonus."
    )
    p2.paragraph_format.space_after = Pt(12)

    # Section 2
    add_heading_1("2. Screen Layout, UI & Visual Hierarchy")
    doc.add_paragraph(
        "The game interface is structured into three primary visual zones:"
    )
    ui_points = [
        ("Top-of-Reels X-Wheels HUD: ", "Displays the 3-tiered wheel progression apparatus: Wheel 1 (Mini), Wheel 2 (Mega), and Wheel 3 (Ultra)."),
        ("5×5 Main Grid Matrix: ", "Contains 25 cell positions indexed from (0,0) at top-left to (4,4) at bottom-right. The central position (2,2) is permanently held by the Central Wild Star in the base game."),
        ("Symbol Life Indicators: ", "Every active coin on the reels features an animated visual life badge (e.g. glowing gems or numeric countdown: 3 -> 2 -> 1 -> Pop)."),
        ("Slingo Payline Overlays: ", "12 predefined winning lines (5 Horizontal Rows, 5 Vertical Columns, and 2 Diagonals).")
    ]
    for bold_txt, norm_txt in ui_points:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r2 = p.add_run(norm_txt)
        p.paragraph_format.space_after = Pt(3)

    # Section 3
    add_heading_1("3. Symbol Catalog & Feature Definitions")
    sym_table = doc.add_table(rows=1, cols=4)
    sym_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = sym_table.rows[0]
    hdr_titles = ["Symbol Name", "Visual Identifier", "Base Cash Value", "Special Function & Gameplay Behavior"]
    for idx, text in enumerate(hdr_titles):
        cell = hdr.cells[idx]
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "1A237E")
        set_cell_margins(cell, 80, 80, 100, 100)

    symbols_data = [
        ("Central Wild Star", "Gold Glowing Star at (2,2)", "0.0x (No cash)", "Permanent Wild in base game. Completes any row, column, or diagonal crossing the center. Never expires and cannot be destroyed."),
        ("Blank", "Dark / Transparent Cell", "0.0x", "Empty space where newly spun symbols can land."),
        ("Cash Coin", "Bronze/Silver/Gold Coin", "0.2x – 5.0x Bet", "Standard cash prize coin. Starts with 3 Lives."),
        ("Jackpot Coin", "Ruby / Sapphire / Diamond", "Mini (5x), Mega (50x), Ultra (500x)", "Fixed Jackpot coin with 3 Lives. Strictly isolated from modifiers (never modified or collected)."),
        ("Mini Strike", "Blue Lightning Coin", "0.2x – 5.0x Bet", "On landing, adds its cash value to all 4 orthogonal neighboring cells (Up, Down, Left, Right)."),
        ("Mega Strike", "Purple Lightning Coin", "0.2x – 5.0x Bet", "On landing, adds its cash value to all symbols sharing any Slingo line with this cell."),
        ("Ultra Strike", "Gold Lightning Coin", "0.2x – 5.0x Bet", "On landing, adds its cash value to all valuable symbols across the entire 5x5 grid."),
        ("Mini Vortex", "Blue Swirling Portal", "Starts at 0.0x", "On landing, gathers and sums cash values from all 4 orthogonal neighbors into itself."),
        ("Mega Vortex", "Purple Swirling Portal", "Starts at 0.0x", "On landing, gathers and sums cash values from all line-sharing symbols into itself."),
        ("Ultra Vortex", "Gold Swirling Portal", "Starts at 0.0x", "On landing, gathers and sums cash values from all coins on the entire grid into itself."),
        ("X Symbol", "Neon Multiplier 'X' Coin", "1.0x Bet", "Starts with 3 Lives. On landing, immediately triggers the Top X-Wheels feature.")
    ]
    for row_idx, data in enumerate(symbols_data):
        row = sym_table.add_row()
        bg_col = "FFFFFF" if row_idx % 2 == 0 else "F9F9F9"
        for col_idx, txt in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = txt
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)
            if col_idx == 0:
                cell.paragraphs[0].runs[0].font.bold = True
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, 60, 60, 80, 80)

    # Section 4
    add_heading_1("4. Base Game Mechanics & Execution Sequence")
    doc.add_paragraph(
        "When the player presses the SPIN button, frontend animations and client state transitions MUST execute in the following exact chronological sequence:"
    )
    steps = [
        ("Step 1: Cleanup & Lifespan Decrement Phase",
         "1. Remove Won Symbols: Any coin that was part of a winning Slingo line on the previous spin fades/pops and frees its space.\n"
         "2. Remove Expired Symbols: Any non-winning coin that had only 1 Life remaining expires and is removed.\n"
         "3. Decrement Surviving Coins: All surviving coins have their life counter reduced by 1 (3 -> 2, or 2 -> 1).\n"
         "4. Note: The Central Wild Star at (2,2) is permanent and is never decremented or cleared."),
        ("Step 2: Symbol Landing Phase",
         "1. New symbols land only on available Blank grid positions.\n"
         "2. Every newly landed symbol initializes with 3 Lives.\n"
         "3. Guaranteed Symbol Rule: At least 1 symbol is guaranteed to land on every spin as long as an empty space exists."),
        ("Step 3: Special Symbol Execution Phase",
         "1. Strikes Animate First: Mini/Mega/Ultra strikes fire lightning animations and add their cash value to target coins.\n"
         "2. Vortexes Animate Second: Mini/Mega/Ultra vortexes pull particle streams and sum target coin values into themselves (original coins retain their values)."),
        ("Step 4: X-Wheel Feature Phase",
         "If an X Symbol landed on the reels, the camera focuses on the Top X-Wheels HUD and spins Wheel 1 (Mini). Landing on Upgrade advances to Wheel 2 (Mega) and potentially Wheel 3 (Ultra). Awards multipliers, ultra strikes, direct jackpots, or Lock & Slingo."),
        ("Step 5: Symbol Life Cycle Reset Phase",
         "Any existing symbol on the reels sharing any Slingo line with any newly landed symbol has its lifespan reset back to 3 Lives!"),
        ("Step 6: 12 Slingo Lines Evaluation Phase",
         "1. A line completes when all 5 positions contain non-blank symbols.\n"
         "2. Line Payout: Player receives the sum of all cash values along that line.\n"
         "3. Intersection Rule: Coins belonging to multiple winning lines pay out for each completed line.\n"
         "4. Winning coins are highlighted and marked to pop at the start of next spin."),
        ("Step 7: Center Wild Wheel Bonus Trigger Phase",
         "If any completed Slingo line crosses through the Central Wild Star at (2,2) (Center Row, Center Column, Main Diagonal, or Anti-Diagonal), the Center Wild Wheel Bonus is triggered once.")
    ]
    for step_title, step_body in steps:
        add_heading_2(step_title)
        p = doc.add_paragraph(step_body)
        p.paragraph_format.space_after = Pt(8)

    # Section 5
    add_heading_1("5. Center Wild Wheel Bonus")
    doc.add_paragraph(
        "Triggered when a completed winning Slingo line crosses the central wild star at (2,2). "
        "An ornate Bonus Wheel appears as an overlay in the center of the reels and spins once to award:"
    )
    wheel_slices = [
        ("Instant Cash Multipliers (1x, 2x, 3x, 4x, 5x): ", "Directly pays 1x to 5x player's total bet."),
        ("Mini Jackpot: ", "Awards 5x total bet."),
        ("Mega Jackpot: ", "Awards 50x total bet."),
        ("Ultra Jackpot: ", "Awards 500x total bet."),
        ("Lock & Slingo: ", "Launches the 5×5 Lock & Slingo™ Bonus Game!")
    ]
    for bold_txt, norm_txt in wheel_slices:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r2 = p.add_run(norm_txt)
        p.paragraph_format.space_after = Pt(3)

    # Section 6
    add_heading_1("6. Lock & Slingo™ Bonus Game")
    doc.add_paragraph(
        "The Lock & Slingo™ Bonus Game is a 5×5 persistent Lock & Win feature with cascading respins:"
    )
    bonus_rules = [
        ("25 Empty Starting Spaces: ", "The bonus begins with an empty 5×5 board. The Central Wild Star does NOT exist in the bonus round (position (2,2) is a standard empty space)."),
        ("Player Lives (3 Respins): ", "Player begins with 3 Lives. If >=1 symbol lands on a spin, lives reset to 3. A blank spin decrements lives by 1."),
        ("Permanent Symbol Locking: ", "Symbols in the bonus never expire. All landed coins remain permanently locked until the bonus ends."),
        ("Active In-Bonus Modifiers: ", "Strikes boost locked coins, Vortexes gather locked coins, and X Symbols spin the bonus X-Wheels."),
        ("Bonus End Conditions: ", "Ends on 3 consecutive blanks (Lives = 0) OR Full House (all 25 positions filled with locked coins).")
    ]
    for bold_txt, norm_txt in bonus_rules:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r2 = p.add_run(norm_txt)
        p.paragraph_format.space_after = Pt(3)

    add_heading_2("Slingo Pay Ladder Awards (Evaluated at End of Bonus)")
    ladder_table = doc.add_table(rows=1, cols=2)
    ladder_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    l_hdr = ladder_table.rows[0]
    l_hdr.cells[0].text = "Completed Slingo Lines"
    l_hdr.cells[1].text = "Highest Achieved Ladder Prize Awarded"
    for c in l_hdr.cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(c, "00796B")
        set_cell_margins(c, 80, 80, 100, 100)

    ladder_data = [
        ("0 – 3 Slingos", "No ladder prize"),
        ("4 Slingos", "Mini Jackpot (5x Total Bet)"),
        ("5 Slingos", "Ultra Strike +1x Boost (Adds +1x to all non-jackpot locked coins)"),
        ("6 Slingos", "Ultra Strike +2x Boost (Adds +2x to all non-jackpot locked coins)"),
        ("7 Slingos", "Ultra Strike +3x Boost (Adds +3x to all non-jackpot locked coins)"),
        ("8 Slingos", "Mega Jackpot (50x Total Bet)"),
        ("9 Slingos", "Ultra Strike +5x Boost (Adds +5x to all non-jackpot locked coins)"),
        ("10 Slingos", "Multiplier x2 Boost (Doubles all non-jackpot locked coins)"),
        ("11 Slingos", "Skipped (Geometry rule - mathematically impossible on 5x5 grid)"),
        ("12 Slingos (Full House)", "Ultra Jackpot (500x Total Bet)")
    ]
    for row_idx, (k, v) in enumerate(ladder_data):
        row = ladder_table.add_row()
        bg_col = "FFFFFF" if row_idx % 2 == 0 else "F4FBF9"
        row.cells[0].text = k
        row.cells[1].text = v
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_background(c, bg_col)
            set_cell_margins(c, 60, 60, 80, 80)
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    p_payout = doc.add_paragraph()
    p_payout.paragraph_format.space_before = Pt(8)
    p_payout.add_run("Final Bonus Payout Formula: ").font.bold = True
    p_payout.add_run("Total Win = Sum of all Locked Grid Cash Values + Highest Achieved Slingo Ladder Prize + Direct Wheel Jackpots.")

    # Section 7
    add_heading_1("7. Critical Isolation Rules & Strict Edge Cases")
    edge_cases = [
        ("A. The Jackpot Isolation Rule (CRITICAL): ", 
         "Jackpot Coins (Mini, Mega, Ultra) are 100% immune to game modifiers. Strikes do not add cash to Jackpot coins, Vortexes do not collect Jackpot coins, and Wheel Multipliers do not multiply Jackpot coins."),
        ("B. The Central Wild Star Isolation Rule: ",
         "The star at (2,2) has 0.0 cash value, never expires, cannot be destroyed or multiplied, is not collected as cash by vortexes, but acts as a wild completing lines."),
        ("C. The 11-Slingo Geometric Skip Rule: ",
         "Filling 24/25 spaces creates 10 lines. Marking the 25th space simultaneously completes both its row and column, jumping the line count from 10 directly to 12. 11 lines can never mathematically exist."),
        ("D. Multi-Line Coin Stacking: ",
         "Coins belonging to intersecting winning lines pay out for each line they belong to."),
        ("E. Single Center Trigger: ",
         "Multiple center-crossing lines on a single spin activate the Wheel Bonus exactly once.")
    ]
    for bold_txt, norm_txt in edge_cases:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r2 = p.add_run(norm_txt)
        p.paragraph_format.space_after = Pt(4)

    # Save Document
    doc.save(output_path)
    print(f"Successfully generated Google Docs / Word specification at: {output_path}")

if __name__ == "__main__":
    output_docx = "/Users/bo.wang/Downloads/Cash_Vortex_Triple_Power_Game_Spec.docx"
    create_game_spec_docx(output_docx)
